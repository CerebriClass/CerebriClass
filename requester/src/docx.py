import math
import random
import string
from datetime import datetime
from typing import List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.shared import Cm, Inches, Mm, Pt, RGBColor
from docx.text.run import Run
from src.types import ResultType


def add_bold_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    return paragraph


def add_bold_paragraph_only_target(document, text, target):
    paragraph = document.add_paragraph()
    runs = paragraph.add_run(text)
    # TODO: 매칭되는 문자열 bold 처리
    # start = 0
    # while True:
    #     start = text.find(target, start)
    #     if start == -1:
    #         break
    #     end = start + len(target)
    #     for i, char in enumerate(runs.text[start:end]):
    #         r_idx = start + i
    #         run = runs.runs[r_idx]
    #         run.bold = True
    #     start = end
    return paragraph


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def generate_random_str(n: int = 10) -> str:
    return ''.join(random.choices(string.ascii_lowercase + string.digits, k=n))


def generate_docx_url():
    return '/downloads/' + datetime.today().strftime('%Y-%m-%d') + '-' + generate_random_str() + '.docx'


CHOICES_CHAR = '①②③④⑤'


def make_dictionary(doc: Document, result: ResultType):
    table = doc.add_table(
        rows=len(result['dictionary']) + 1, cols=2, style=doc.styles["Table Grid"])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for y in range(len(table.rows)):
        row = table.rows[y]
        if y == 0:
            row.cells[0].text = 'Vocabulary'
            row.cells[1].text = '뜻'
            continue
        for x in range(len(row.cells)):
            cell = row.cells[x]
            if x == 0:
                cell.text = result['dictionary'][y - 1]['word']
            else:
                cell.text = result['dictionary'][y - 1]['meaning']

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.autofit = False
    table.allow_autofit = False
    set_col_widths(table, (Cm(3), Cm(3)))


def make_example_sentences(doc: Document, result: ResultType):
    add_bold_paragraph(doc, '[예문]')
    for i, elem in enumerate(result['example_sentences']):
        add_bold_paragraph(doc, f"{i+1}. {elem['word']}")
        for sentence in elem['sentences']:
            add_bold_paragraph_only_target(
                doc, sentence['sentence'], elem['word'])
            doc.add_paragraph(f"{sentence['meaning']}")


def make_eng_word_quiz(doc: Document, result: ResultType, no: int, answer: List[int]):
    add_bold_paragraph(doc, '[뜻 선택]')
    add_bold_paragraph(doc, '다음 중 주어진 단어의 뜻으로 가장 적절한 보기를 고르시오.')

    for elem in result['eng_word_quiz']:
        add_bold_paragraph(doc, f"{no}. {elem['word']}")
        doc.add_paragraph(
            '\n'.join([f"  {CHOICES_CHAR[i]} {choice}" for i, choice in enumerate(elem['choices'])]))
        answer.append({'no': no, 'answer': elem['answer'] + 1})
        no += 1
    return no, answer


def make_kor_word_quiz(doc: Document, result: ResultType, no: int, answer: List[int]):
    add_bold_paragraph(doc, '[단어 선택]')
    add_bold_paragraph(doc, '다음 중 해당 뜻을 가진 단어로 가장 적절한 보기를 고르시오.')
    for elem in result['kor_word_quiz']:
        add_bold_paragraph(doc, f"{no}. {elem['word']}")
        doc.add_paragraph(
            '\n'.join([f"  {CHOICES_CHAR[i]} {choice}" for i, choice in enumerate(elem['choices'])]))
        answer.append({'no': no, 'answer': elem['answer'] + 1})
        no += 1
    return no, answer


def make_blank_quiz(doc: Document, result: ResultType, no: int, answer: List[int]):
    add_bold_paragraph(doc, '[단어 선택]')
    add_bold_paragraph(doc, '빈칸에 들어갈 단어 중 가장 적절한 것을 고르시오.')
    for elem in result['blank_quiz']:
        add_bold_paragraph(doc, f"{no}. {elem['sentence']}")
        doc.add_paragraph(
            '\n'.join([f"  {CHOICES_CHAR[i]} {choice}" for i, choice in enumerate(elem['choices'])]))
        answer.append({'no': no, 'answer': elem['answer'] + 1})
        no += 1
    return no, answer


def make_doc(result: ResultType):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Mm(19.1)
    section.right_margin = Mm(19.1)
    # multiple columns
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    normal_style = doc.styles['Normal']
    # style._element.rPr.rFonts.set(qn('w:eastAsia'), 'HY신명조') # FIXME: 한글 폰트 적용 안됨
    normal_style.font.name = 'HY신명조'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'HY신명조')
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.line_spacing = 1.2

    h1_style = doc.styles['Heading 1']
    h1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2_style = doc.styles['Heading 2']
    h2_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h3_style = doc.styles['Heading 3']
    h3_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    h1 = doc.add_heading('영어 단어 학습지', 1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('Storyboard.lab')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('이름: __________________')
    doc.add_paragraph('날짜: __________________')
    doc.add_paragraph()

    make_dictionary(doc, result)
    doc.add_paragraph()

    make_example_sentences(doc, result)
    doc.add_paragraph()

    no = 1
    answer = []
    if len(result['eng_word_quiz']):
        no, answer = make_eng_word_quiz(doc, result, no, answer)
        doc.add_paragraph()
    if len(result['kor_word_quiz']):
        no, answer = make_kor_word_quiz(doc, result, no, answer)
        doc.add_paragraph()
    if len(result['blank_quiz']):
        no, answer = make_blank_quiz(doc, result, no, answer)
        doc.add_paragraph()

    doc.add_page_break()
    # fast answer
    doc.add_heading('[빠른 정답]', 2)
    table = doc.add_table(rows=math.ceil(len(answer) / 5), cols=10,
                          style=doc.styles["Table Grid"])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(0, len(answer)):
        y, x = i // 5, i % 5
        row = doc.tables[1].rows[y]
        row.cells[x * 2].text = str(answer[i]['no'])
        row.cells[x * 2 + 1].text = CHOICES_CHAR[answer[i]['answer'] - 1]

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return doc


def make_sheet(result: ResultType):
    doc = make_doc(result)
    url = generate_docx_url()
    doc.save(url)
    return url
